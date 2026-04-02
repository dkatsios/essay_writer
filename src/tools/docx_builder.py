"""DOCX document assembly.

Converts structured essay text (markdown-like) into a formatted .docx file
with cover page, table of contents, headings, body text, and references.
"""

from __future__ import annotations

import json
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def _safe_json_loads(s: str) -> dict:
    """Parse JSON, handling double-escaped strings from LLM output."""
    try:
        return json.loads(s)
    except json.JSONDecodeError:
        return json.loads(s.encode().decode("unicode_escape"))


def _set_document_defaults(doc: Document, config: dict) -> None:
    """Apply default formatting to the document."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = config.get("font", "Times New Roman")
    font.size = Pt(config.get("font_size", 12))

    pf = style.paragraph_format
    pf.space_after = Pt(0)

    line_spacing = config.get("line_spacing", 1.5)
    pf.line_spacing = line_spacing

    alignment = config.get("text_alignment", "justified")
    alignment_map = {
        "justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    pf.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    if config.get("paragraph_indent", True):
        pf.first_line_indent = Cm(1.27)

    for section in doc.sections:
        margin = Cm(config.get("margins_cm", 2.5))
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin


def _add_cover_page(doc: Document, config: dict) -> None:
    """Add a cover page with title, author, institution, course, date."""
    # Add some blank space
    for _ in range(6):
        doc.add_paragraph()

    title = config.get("title", "")
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(18)

    doc.add_paragraph()

    for field in ["author", "institution", "course", "professor", "date"]:
        value = config.get(field, "")
        if value:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(value).font.size = Pt(14)

    doc.add_page_break()


def _normalize_toc_styles(doc: Document, font_name: str, font_size: Pt) -> None:
    """Ensure TOC heading styles use the document font instead of template defaults."""
    for level in range(1, 5):
        style_name = f"TOC {level}"
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = font_name
            style.font.size = font_size
            style.paragraph_format.space_after = Pt(4)


def _add_native_toc(doc: Document) -> None:
    """Insert a native Word TOC field that auto-updates when opened."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("ΠΙΝΑΚΑΣ ΠΕΡΙΕΧΟΜΕΝΩΝ")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()

    # Insert TOC field code
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)

    run1 = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run1._element.append(fld_begin)

    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-4" \\h \\z \\u '
    run2._element.append(instr)

    run3 = p.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_sep)

    run4 = p.add_run("[Ενημερώστε τον πίνακα περιεχομένων]")
    run4.font.size = Pt(10)

    run5 = p.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_end)

    doc.add_page_break()

    # Tell Word to update all fields (including the TOC) on open
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    doc.settings.element.append(update_fields)


def _add_page_numbers(doc: Document, position: str = "bottom_center") -> None:
    """Add page numbers to the document footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.text = "PAGE"
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run = p.add_run()
        run._element.append(fld_char_begin)
        run2 = p.add_run()
        run2._element.append(instr_text)
        run3 = p.add_run()
        run3._element.append(fld_char_end)


_HEADING_RE = re.compile(r"^(#{1,4})\s+(.+)$")
_BULLET_RE = re.compile(r"^[\*\-]\s+(.+)$")
_NUMBERED_RE = re.compile(r"^\d+\.\s+(.+)$")
_INLINE_RE = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|\^\^(\d+)\^\^)")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|\s*$")
_TABLE_SEP_RE = re.compile(r"^\|[\s:]*-{3,}[\s:]*")
_MD_BOLD_RE = re.compile(r"\*{1,3}(.+?)\*{1,3}")

# Citation markers: [[source_id]] or [[source_id|σ. 15]]
# Normalize [[a], [b], [c]] → [[a]] [[b]] [[c]] (common LLM mistake)
_MULTI_CITE_RE = re.compile(r"\[\[([^\]|]+?)(?:\],\s*\[([^\]|]+?))+\]\]")


def _normalize_citations(text: str) -> str:
    """Fix grouped citation markers into separate ones."""

    def _split(m: re.Match) -> str:
        inner = m.group(0)[2:-2]  # strip outer [[ and ]]
        ids = re.split(r"\],\s*\[", inner)
        return " ".join(f"[[{cid.strip()}]]" for cid in ids)

    return _MULTI_CITE_RE.sub(_split, text)


_CITE_RE = re.compile(r"\[\[([^\]|]+?)(?:\|([^\]]+?))?\]\]")


def _add_formatted_runs(paragraph, text: str) -> None:
    """Parse markdown inline formatting and add runs with bold/italic/superscript."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        if m.group(2):  # ***bold+italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(3):  # **bold**
            run = paragraph.add_run(m.group(3))
            run.bold = True
        elif m.group(4):  # *italic*
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(5):  # ^^N^^ superscript (footnote number)
            run = paragraph.add_run(m.group(5))
            run.font.superscript = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------------------------------------------------------------------------
# Citation processing
# ---------------------------------------------------------------------------


def _has_valid_authors(source: dict) -> bool:
    """Return True if the source has at least one non-blank author."""
    return any(a.strip() for a in source.get("authors", []))


def _clean_authors(source: dict) -> list[str]:
    """Return author list with blank entries removed."""
    return [a for a in source.get("authors", []) if a.strip()]


def _format_apa_inline(source: dict, page_info: str | None) -> str:
    """Format an APA7 in-text citation string."""
    authors = _clean_authors(source)
    year = source.get("year", "n.d.")
    if not authors:
        name = source.get("title", "Unknown")[:30]
    elif len(authors) == 1:
        name = authors[0].split(",")[0].strip()
    elif len(authors) == 2:
        parts = [a.split(",")[0].strip() for a in authors]
        name = f"{parts[0]} & {parts[1]}"
    else:
        name = f"{authors[0].split(',')[0].strip()} et al."
    citation = f"({name}, {year}"
    if page_info:
        citation += f", {page_info}"
    citation += ")"
    return citation


def _format_bib_entry(source: dict) -> str:
    """Format a full APA7 bibliography entry."""
    authors = _clean_authors(source)
    year = source.get("year", "n.d.")
    title = source.get("title", "")
    journal = source.get("source", "")
    volume = source.get("volume", "")
    issue = source.get("issue", "")
    pages = source.get("pages", "")
    doi = source.get("doi", "")
    url = source.get("url", "")
    publisher = source.get("publisher", "")

    author_str = ", ".join(authors) if authors else "Unknown"
    entry = f"{author_str} ({year}). {title}."
    if journal:
        entry += f" *{journal}*"
        if volume:
            entry += f", *{volume}*"
            if issue:
                entry += f"({issue})"
        if pages:
            entry += f", {pages}"
        entry += "."
    elif publisher:
        entry += f" {publisher}."
    if doi:
        entry += f" https://doi.org/{doi}"
    elif url:
        entry += f" {url}"
    return entry


def _process_citations(essay_text: str, sources: dict, style: str) -> str:
    """Replace [[source_id]] markers with formatted citations.

    For 'apa7': (Author, Year) inline + Βιβλιογραφία section.
    For 'footnotes': superscript ^^N^^ markers + Σημειώσεις section.
    """
    if not sources:
        return essay_text

    essay_text = _normalize_citations(essay_text)

    used: list[tuple[str, str | None]] = []  # (source_id, page_info)
    seen: dict[str, int] = {}

    def replacer(m: re.Match) -> str:
        source_id = m.group(1).strip()
        page_info = m.group(2).strip() if m.group(2) else None
        source = sources.get(source_id)
        if not source:
            return f"[{source_id}]"

        if style == "footnotes":
            used.append((source_id, page_info))
            return f"^^{len(used)}^^"
        else:  # apa7
            if source_id not in seen:
                seen[source_id] = True
                used.append((source_id, page_info))
            return _format_apa_inline(source, page_info)

    processed = _CITE_RE.sub(replacer, essay_text)

    if not used:
        return processed

    if style == "footnotes":
        processed += "\n\n## Σημειώσεις\n\n"
        for i, (sid, page_info) in enumerate(used, 1):
            source = sources.get(sid)
            if source:
                entry = _format_bib_entry(source)
                if page_info:
                    entry += f" [{page_info}]"
                processed += f"{i}. {entry}\n\n"
    else:  # apa7
        processed += "\n\n## Βιβλιογραφία\n\n"
        unique_ids = list(dict.fromkeys(sid for sid, _ in used))
        entries = []
        for sid in unique_ids:
            source = sources.get(sid)
            if source:
                entries.append(_format_bib_entry(source))
        entries.sort()
        processed += "\n\n".join(entries)

    return processed


def _parse_table_rows(lines: list[str]) -> list[list[str]]:
    """Parse pipe-delimited markdown table rows into a list of cell lists."""
    rows: list[list[str]] = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    return rows


def _add_table(
    doc: Document, header_cells: list[str], data_rows: list[list[str]]
) -> None:
    """Add a formatted table to the document."""
    num_cols = len(header_cells)
    table = doc.add_table(rows=1, cols=num_cols, style="Table Grid")

    # Header row
    for i, cell_text in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(cell_text)
        run.bold = True

    # Data rows
    for row_cells in data_rows:
        row = table.add_row()
        for i, cell_text in enumerate(row_cells[:num_cols]):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_formatted_runs(p, cell_text)

    # Add a blank paragraph after the table for spacing
    doc.add_paragraph()


def _parse_and_add_content(doc: Document, essay_text: str) -> None:
    """Parse markdown-like essay text and add to document with proper styles."""
    lines = essay_text.split("\n")
    current_paragraph_lines: list[str] = []
    skipped_first_h1 = False
    in_bibliography = False

    def flush_paragraph() -> None:
        if current_paragraph_lines:
            text = " ".join(current_paragraph_lines)
            p = doc.add_paragraph()
            _add_formatted_runs(p, text)
            if in_bibliography:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_after = Pt(6)
            current_paragraph_lines.clear()

    i = 0
    while i < len(lines):
        stripped = lines[i].strip()

        # Detect markdown table: header row, separator row, data rows
        if (
            _TABLE_ROW_RE.match(stripped)
            and i + 1 < len(lines)
            and _TABLE_SEP_RE.match(lines[i + 1].strip())
        ):
            flush_paragraph()
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip header and separator
            data_lines: list[str] = []
            while i < len(lines) and _TABLE_ROW_RE.match(lines[i].strip()):
                data_lines.append(lines[i])
                i += 1
            data_rows = _parse_table_rows(data_lines)
            _add_table(doc, header_cells, data_rows)
            continue

        heading_match = _HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = _MD_BOLD_RE.sub(r"\1", heading_match.group(2)).strip()
            # Skip the first H1 — it's the essay title already on the cover page
            if level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
                i += 1
                continue
            flush_paragraph()
            # Detect bibliography/references section for left-aligned formatting
            lower_text = text.lower()
            if lower_text in (
                "βιβλιογραφία",
                "σημειώσεις",
                "references",
                "bibliography",
            ):
                in_bibliography = True
            doc.add_heading(text, level=min(level, 4))
        elif stripped == "":
            flush_paragraph()
        elif _BULLET_RE.match(stripped):
            flush_paragraph()
            content = _BULLET_RE.match(stripped).group(1)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, content)
        elif _NUMBERED_RE.match(stripped):
            flush_paragraph()
            content = _NUMBERED_RE.match(stripped).group(1)
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, content)
        else:
            current_paragraph_lines.append(stripped)

        i += 1

    flush_paragraph()


def build_document(
    essay_text: str, config: dict, sources: dict | None = None
) -> Document:
    """Build a complete .docx document from essay text and config."""
    doc = Document()
    _set_document_defaults(doc, config)
    font_name = config.get("font", "Times New Roman")
    font_size = Pt(config.get("font_size", 12))
    _normalize_toc_styles(doc, font_name, font_size)
    _add_cover_page(doc, config)
    _add_native_toc(doc)

    if sources:
        style = config.get("citation_style", "apa7")
        essay_text = _process_citations(essay_text, sources, style)

    _parse_and_add_content(doc, essay_text)
    _add_page_numbers(doc, config.get("page_numbers", "bottom_center"))
    return doc
