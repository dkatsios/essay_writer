"""Tests for src/intake.py — file classification, image blocks, text extraction."""

from __future__ import annotations

import base64
from pathlib import Path

import pytest


class TestClassify:
    @pytest.mark.parametrize(
        "filename,expected",
        [
            ("doc.pdf", "pdf"),
            ("doc.docx", "docx"),
            ("doc.pptx", "pptx"),
            ("image.png", "image"),
            ("image.jpg", "image"),
            ("notes.txt", "text"),
            ("notes.md", "text"),
            ("data.csv", "text"),
            ("file.xyz", "unsupported"),
        ],
    )
    def test_classify(self, filename, expected):
        from src.intake import classify

        assert classify(Path(filename)) == expected


class TestMakeImageBlock:
    def test_produces_valid_block(self):
        from src.intake import make_image_block

        block = make_image_block(b"fake-png-data", "image/png")
        assert block["type"] == "image_url"
        url = block["image_url"]["url"]
        assert url.startswith("data:image/png;base64,")
        # Verify the base64 decodes back
        encoded = url.split(",", 1)[1]
        assert base64.standard_b64decode(encoded) == b"fake-png-data"


class TestBuildExtractedText:
    def test_builds_extracted_text_with_prompt_and_warnings(self):
        from src.intake import InputFile, build_extracted_text

        files = [
            InputFile(Path("topic.txt"), "text", text="Essay topic"),
            InputFile(Path("scan.pdf"), "pdf", image_blocks=[{"type": "image_url"}]),
            InputFile(
                Path("legacy.doc"),
                "unsupported",
                warning="Old Word binary format — save as .docx first",
            ),
        ]

        extracted = build_extracted_text(files, extra_prompt="Focus on economics")

        assert "### File: topic.txt" in extracted
        assert "Essay topic" in extracted
        assert "### Image: scan.pdf" in extracted
        assert "text extraction was sparse" in extracted
        assert "## Warnings" in extracted
        assert "legacy.doc" in extracted
        assert "## Additional Instructions" in extracted
        assert "Focus on economics" in extracted


class TestExtractDocxText:
    def test_extracts_headings_and_body(self):
        from src.tools.docx_reader import extract_docx_text

        from docx import Document

        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Body text here.")
        doc.add_heading("Sub", level=2)

        result = extract_docx_text(doc)
        assert "# Title" in result
        assert "Body text here." in result
        assert "## Sub" in result
