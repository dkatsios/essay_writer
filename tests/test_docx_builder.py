"""Tests for src/tools/docx_builder.py — tables, headings, citations, alignment."""

from __future__ import annotations


class TestDocxTableParsing:
    """Tests for markdown table -> docx table conversion."""

    def test_simple_table(self):
        from src.tools.docx_builder import build_document

        md = (
            "Some text before.\n"
            "\n"
            "| Name | Age |\n"
            "|------|-----|\n"
            "| Alice | 30 |\n"
            "| Bob | 25 |\n"
            "\n"
            "Some text after."
        )
        doc = build_document(md, {"title": "Test"})
        tables = doc.tables
        assert len(tables) == 1
        table = tables[0]
        # Header row + 2 data rows
        assert len(table.rows) == 3
        assert len(table.columns) == 2
        assert table.rows[0].cells[0].text == "Name"
        assert table.rows[0].cells[1].text == "Age"
        assert table.rows[1].cells[0].text == "Alice"
        assert table.rows[2].cells[1].text == "25"

    def test_table_with_inline_formatting(self):
        from src.tools.docx_builder import build_document

        md = "| Header |\n|--------|\n| **bold** and *italic* |\n"
        doc = build_document(md, {"title": "Test"})
        tables = doc.tables
        assert len(tables) == 1
        # Cell should contain the text (formatting applied via runs)
        cell_text = tables[0].rows[1].cells[0].text
        assert "bold" in cell_text
        assert "italic" in cell_text

    def test_three_column_table(self):
        from src.tools.docx_builder import build_document

        md = (
            "| A | B | C |\n"
            "|---|---|---|\n"
            "| 1 | 2 | 3 |\n"
            "| 4 | 5 | 6 |\n"
            "| 7 | 8 | 9 |\n"
        )
        doc = build_document(md, {"title": "Test"})
        table = doc.tables[0]
        assert len(table.columns) == 3
        assert len(table.rows) == 4  # 1 header + 3 data
        assert table.rows[3].cells[2].text == "9"

    def test_no_table_without_separator(self):
        """Pipe lines without a separator row should NOT be parsed as a table."""
        from src.tools.docx_builder import build_document

        md = "| Not a table |\n| Just pipes |\n"
        doc = build_document(md, {"title": "Test"})
        assert len(doc.tables) == 0

    def test_table_between_paragraphs(self):
        """Table should not swallow surrounding paragraphs."""
        from src.tools.docx_builder import build_document

        md = "Paragraph before.\n\n| X |\n|---|\n| 1 |\n\nParagraph after."
        doc = build_document(md, {"title": "Test"})
        assert len(doc.tables) == 1
        # Check that both paragraphs exist in the document text
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Paragraph before." in full_text
        assert "Paragraph after." in full_text

    def test_header_bold(self):
        """Header cells should be bold."""
        from src.tools.docx_builder import build_document

        md = "| Col |\n|-----|\n| val |\n"
        doc = build_document(md, {"title": "Test"})
        header_cell = doc.tables[0].rows[0].cells[0]
        runs = header_cell.paragraphs[0].runs
        assert any(r.bold for r in runs)


class TestHeadingAsterisks:
    """Headings should have markdown bold/italic markers stripped."""

    def test_strips_double_asterisks(self):
        from src.tools.docx_builder import build_document

        md = "## **Bold Heading**\n\nSome text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert any("Bold Heading" in h and "**" not in h for h in headings)

    def test_strips_single_asterisks(self):
        from src.tools.docx_builder import build_document

        md = "## *Italic Heading*\n\nSome text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert any("Italic Heading" in h and "*" not in h for h in headings)

    def test_strips_triple_asterisks(self):
        from src.tools.docx_builder import build_document

        md = "## ***Bold Italic Heading***\n\nSome text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert any("Bold Italic Heading" in h and "*" not in h for h in headings)


class TestEmptyAuthorCitations:
    """Sources with empty/blank authors should fall back to title-based citation."""

    def test_empty_author_list_inline(self):
        from src.tools.docx_builder import format_apa_inline

        source = {"authors": [], "title": "Some Title", "year": 2020}
        result = format_apa_inline(source, None)
        assert "Some Title" in result
        assert "& " not in result

    def test_blank_authors_inline(self):
        from src.tools.docx_builder import format_apa_inline

        source = {"authors": ["", "  "], "title": "My Paper", "year": 2017}
        result = format_apa_inline(source, None)
        assert "My Paper" in result
        assert "( & " not in result

    def test_blank_authors_bib_entry(self):
        from src.tools.docx_builder import format_bib_entry

        source = {"authors": ["", ""], "title": "A Title", "year": 2020}
        result = format_bib_entry(source)
        assert result.startswith("Unknown (2020)")
        assert ", ," not in result


class TestH1TitleSkipping:
    """First H1 should be skipped since it duplicates the cover page title."""

    def test_first_h1_skipped(self):
        from src.tools.docx_builder import build_document

        md = "# Essay Title\n\n## Section One\n\nSome text."
        doc = build_document(md, {"title": "Essay Title"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert "Essay Title" not in headings
        assert "Section One" in headings

    def test_second_h1_kept(self):
        from src.tools.docx_builder import build_document

        md = "# First Title\n\n## Section\n\nText.\n\n# Second H1\n\nMore text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert "First Title" not in headings
        assert "Second H1" in headings


class TestTableCellAlignment:
    """Table cells should be left-aligned, not justified."""

    def test_cells_left_aligned(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        from src.tools.docx_builder import build_document

        md = "| A | B |\n|---|---|\n| 1 | 2 |\n"
        doc = build_document(md, {"title": "Test"})
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
