"""Tests for the essay writer pipeline modules."""

from __future__ import annotations

import base64
import asyncio
import json
import logging
import threading
import time
from pathlib import Path
from unittest.mock import MagicMock

import httpx
import pytest


def _google_service_account_json(project_id: str = "service-project") -> str:
    return json.dumps(
        {
            "type": "service_account",
            "project_id": project_id,
            "private_key_id": "private-key-id",
            "private_key": "-----BEGIN PRIVATE KEY-----\nabc123\n-----END PRIVATE KEY-----\n",
            "client_email": "essay-writer@test-project.iam.gserviceaccount.com",
            "client_id": "1234567890",
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "auth_provider_x509_cert_url": "https://www.googleapis.com/oauth2/v1/certs",
            "client_x509_cert_url": "https://www.googleapis.com/robot/v1/metadata/x509/test",
        },
        indent=2,
    )


# ── agent retry logic ─────────────────────────────────────────────────


class TestRetryWithBackoff:
    """Tests for _retry_with_backoff in agent.py."""

    def test_immediate_success(self):
        from src.agent import _retry_with_backoff

        result = _retry_with_backoff(lambda: "hello")
        assert result == "hello"

    def test_retries_on_resource_exhausted(self):
        from src.agent import _retry_with_backoff

        calls = []

        def fn():
            calls.append(1)
            if len(calls) == 1:
                raise Exception("429 RESOURCE_EXHAUSTED")
            return "ok"

        import src.agent

        original_sleep = src.agent.time.sleep
        src.agent.time.sleep = lambda _: None
        try:
            result = _retry_with_backoff(fn)
            assert result == "ok"
            assert len(calls) == 2
        finally:
            src.agent.time.sleep = original_sleep


class TestGoogleProviderNormalization:
    def test_google_classic_api_key_stays_on_google_provider(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.setenv("GOOGLE_API_KEY", "AIza-classic-key")
        monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_LOCATION", raising=False)

        model, kwargs = _normalize_model_spec("google_genai:gemini-2.5-flash")

        assert model == "google/gemini-2.5-flash"
        assert kwargs == {"api_key": "AIza-classic-key"}

    def test_google_vertex_api_key_routes_to_vertex_provider(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.setenv("GOOGLE_CLOUD_PROJECT", "demo-project")
        monkeypatch.setenv("GOOGLE_CLOUD_LOCATION", "us-central1")

        model, kwargs = _normalize_model_spec(
            "google_genai:gemini-2.5-flash",
            api_key="AQ.vertex-key",
        )

        assert model == "vertexai/gemini-2.5-flash"
        assert kwargs == {
            "api_key": "AQ.vertex-key",
            "project": "demo-project",
            "location": "us-central1",
        }

    def test_google_vertex_api_key_requires_project_and_location(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_LOCATION", raising=False)

        with pytest.raises(
            ValueError,
            match="GOOGLE_CLOUD_PROJECT and GOOGLE_CLOUD_LOCATION",
        ):
            _normalize_model_spec(
                "google_genai:gemini-2.5-flash",
                api_key="AQ.vertex-key",
            )

    def test_google_service_account_json_routes_to_vertex_provider(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
        monkeypatch.setenv("GOOGLE_CLOUD_LOCATION", "us-central1")

        model, kwargs = _normalize_model_spec(
            "google_genai:gemini-2.5-flash",
            api_key=_google_service_account_json(),
        )

        assert model == "vertexai/gemini-2.5-flash"
        assert kwargs == {
            "project": "service-project",
            "location": "us-central1",
        }

    def test_google_service_account_json_prefers_env_project(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.setenv("GOOGLE_CLOUD_PROJECT", "env-project")
        monkeypatch.setenv("GOOGLE_CLOUD_LOCATION", "us-central1")

        model, kwargs = _normalize_model_spec(
            "google_genai:gemini-2.5-flash",
            api_key=_google_service_account_json(project_id="json-project"),
        )

        assert model == "vertexai/gemini-2.5-flash"
        assert kwargs == {
            "project": "env-project",
            "location": "us-central1",
        }

    def test_google_service_account_json_requires_location(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_LOCATION", raising=False)

        with pytest.raises(
            ValueError,
            match="GOOGLE_CLOUD_LOCATION and either GOOGLE_CLOUD_PROJECT or a service-account JSON project_id",
        ):
            _normalize_model_spec(
                "google_genai:gemini-2.5-flash",
                api_key=_google_service_account_json(),
            )

    def test_google_service_account_json_parse_error_is_explicit(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)

        with pytest.raises(
            ValueError,
            match="looks like JSON but could not be parsed",
        ):
            _normalize_model_spec(
                "google_genai:gemini-2.5-flash",
                api_key='{"type": "service_account",',
            )

    def test_explicit_google_vertexai_provider_uses_vertex_metadata(self, monkeypatch):
        from src.agent import _normalize_model_spec

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.setenv("GOOGLE_CLOUD_PROJECT", "demo-project")
        monkeypatch.setenv("GOOGLE_CLOUD_LOCATION", "europe-west4")
        monkeypatch.setenv("GOOGLE_API_KEY", "AQ.vertex-key")

        model, kwargs = _normalize_model_spec("google_vertexai:gemini-2.5-flash")

        assert model == "vertexai/gemini-2.5-flash"
        assert kwargs == {
            "api_key": "AQ.vertex-key",
            "project": "demo-project",
            "location": "europe-west4",
        }

    def test_gateway_mode_warns_when_vertex_google_key_is_present(
        self, monkeypatch, caplog
    ):
        from src.agent import _normalize_model_spec

        monkeypatch.setenv("AI_BASE_URL", "https://gateway.example.com")
        monkeypatch.setenv("AI_API_KEY", "gateway-key")
        monkeypatch.setenv("GOOGLE_API_KEY", "AQ.vertex-key")
        monkeypatch.delenv("AI_MODEL", raising=False)

        with caplog.at_level(logging.WARNING):
            model, kwargs = _normalize_model_spec("google_genai:gemini-2.5-flash")

        assert model == "openai/vertex_ai.gemini-2.5-flash"
        assert kwargs == {
            "base_url": "https://gateway.example.com",
            "api_key": "gateway-key",
        }
        assert "direct Google Vertex credentials are skipped" in caplog.text

    def test_gateway_mode_warns_when_google_service_account_json_is_present(
        self, monkeypatch, caplog
    ):
        from src.agent import _normalize_model_spec

        monkeypatch.setenv("AI_BASE_URL", "https://gateway.example.com")
        monkeypatch.setenv("AI_API_KEY", "gateway-key")
        monkeypatch.delenv("AI_MODEL", raising=False)

        with caplog.at_level(logging.WARNING):
            model, kwargs = _normalize_model_spec(
                "google_genai:gemini-2.5-flash",
                api_key=_google_service_account_json(),
            )

        assert model == "openai/vertex_ai.gemini-2.5-flash"
        assert kwargs == {
            "base_url": "https://gateway.example.com",
            "api_key": "gateway-key",
        }
        assert "direct Google Vertex credentials are skipped" in caplog.text

    def test_create_client_uses_from_genai_for_service_account(self, monkeypatch):
        from src.agent import create_client

        captured: dict[str, object] = {}

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
        monkeypatch.setenv("GOOGLE_CLOUD_LOCATION", "us-central1")

        def fake_from_service_account_info(info, scopes):
            captured["service_account_info"] = info
            captured["scopes"] = scopes
            return "credentials"

        def fake_genai_client(**kwargs):
            captured["genai_client_kwargs"] = kwargs
            return "raw-client"

        def fake_from_genai(raw_client, use_async=False, **kwargs):
            captured["from_genai"] = {
                "raw_client": raw_client,
                "use_async": use_async,
                "kwargs": kwargs,
            }
            return "wrapped-client"

        monkeypatch.setattr(
            "google.oauth2.service_account.Credentials.from_service_account_info",
            fake_from_service_account_info,
        )
        monkeypatch.setattr("google.genai.Client", fake_genai_client)
        monkeypatch.setattr("src.agent.instructor.from_genai", fake_from_genai)
        monkeypatch.setattr(
            "src.agent.instructor.from_provider",
            lambda *_args, **_kwargs: pytest.fail("from_provider should not be used"),
        )

        client = create_client(
            "google_genai:gemini-2.5-flash",
            api_key=_google_service_account_json(),
        )

        assert client.client == "wrapped-client"
        assert client.model == "gemini-2.5-flash"
        assert captured["scopes"] == ["https://www.googleapis.com/auth/cloud-platform"]
        assert captured["genai_client_kwargs"] == {
            "vertexai": True,
            "credentials": "credentials",
            "project": "service-project",
            "location": "us-central1",
        }
        assert captured["from_genai"] == {
            "raw_client": "raw-client",
            "use_async": False,
            "kwargs": {},
        }

    def test_create_async_client_uses_from_genai_for_service_account(self, monkeypatch):
        from src.agent import create_async_client

        captured: dict[str, object] = {}

        monkeypatch.delenv("AI_BASE_URL", raising=False)
        monkeypatch.delenv("GOOGLE_CLOUD_PROJECT", raising=False)
        monkeypatch.setenv("GOOGLE_CLOUD_LOCATION", "us-central1")

        monkeypatch.setattr(
            "google.oauth2.service_account.Credentials.from_service_account_info",
            lambda info, scopes: (
                captured.update({"service_account_info": info, "scopes": scopes})
                or "credentials"
            ),
        )
        monkeypatch.setattr(
            "google.genai.Client",
            lambda **kwargs: (
                captured.update({"genai_client_kwargs": kwargs}) or "raw-client"
            ),
        )
        monkeypatch.setattr(
            "src.agent.instructor.from_genai",
            lambda raw_client, use_async=False, **kwargs: (
                captured.update(
                    {
                        "from_genai": {
                            "raw_client": raw_client,
                            "use_async": use_async,
                            "kwargs": kwargs,
                        }
                    }
                )
                or "wrapped-client"
            ),
        )
        monkeypatch.setattr(
            "src.agent.instructor.from_provider",
            lambda *_args, **_kwargs: pytest.fail("from_provider should not be used"),
        )

        client = create_async_client(
            "google_genai:gemini-2.5-flash",
            api_key=_google_service_account_json(),
        )

        assert client.client == "wrapped-client"
        assert client.model == "gemini-2.5-flash"
        assert captured["from_genai"] == {
            "raw_client": "raw-client",
            "use_async": True,
            "kwargs": {},
        }

    def test_retries_on_timeout(self):
        from src.agent import _retry_with_backoff

        calls = []

        def fn():
            calls.append(1)
            if len(calls) == 1:
                raise TimeoutError("Request timed out")
            return "ok"

        import src.agent

        original_sleep = src.agent.time.sleep
        src.agent.time.sleep = lambda _: None
        try:
            result = _retry_with_backoff(fn)
            assert result == "ok"
            assert len(calls) == 2
        finally:
            src.agent.time.sleep = original_sleep


def test_model_client_to_async_does_not_store_api_key(monkeypatch):
    from src.agent import ModelClient

    captured = {}

    def fake_create_async_client(model_spec, *, api_key=None):
        captured["model_spec"] = model_spec
        captured["api_key"] = api_key
        return "async-client"

    monkeypatch.setattr("src.agent.create_async_client", fake_create_async_client)

    client = ModelClient(
        client=object(),
        model="gpt-5.4",
        model_spec="openai:gpt-5.4",
    )

    result = client.to_async()

    assert result == "async-client"
    assert captured == {
        "model_spec": "openai:gpt-5.4",
        "api_key": None,
    }


class TestStructuredCallRepair:
    def test_essay_plan_parses_stringified_sections(self):
        from src.schemas import EssayPlan

        plan = EssayPlan.model_validate(
            {
                "title": "Test",
                "thesis": "Test thesis",
                "sections": json.dumps(
                    [
                        {
                            "number": 1,
                            "title": "Intro",
                            "heading": "Intro",
                            "word_target": 400,
                        },
                        {
                            "number": 2,
                            "title": "Body",
                            "heading": "Body",
                            "word_target": 600,
                        },
                    ]
                ),
                "research_queries": ["query"],
                "total_word_target": 1000,
            }
        )

        assert len(plan.sections) == 2
        assert plan.sections[0].title == "Intro"

    def test_essay_plan_rejects_missing_sections(self):
        from src.schemas import EssayPlan

        with pytest.raises(ValueError, match="sections must be a non-empty array"):
            EssayPlan(
                title="Η πτώση της εμπιστοσύνης",
                thesis="Η ανάλυση απαιτεί σύνθεση κοινωνικών και θεσμικών παραγόντων.",
                research_queries=["πτώση εμπιστοσύνης θεσμοί Ελλάδα"],
                total_word_target=1200,
            )

    def test_structured_call_uses_instructor(self, monkeypatch):
        """Verify _structured_call delegates to Instructor's create()."""
        from src.pipeline_support import _structured_call
        from src.schemas import EssayPlan
        from src.agent import ModelClient

        complete_plan = EssayPlan.model_validate(
            {
                "title": "Test",
                "thesis": "Test thesis",
                "research_queries": ["query"],
                "total_word_target": 1000,
                "sections": [
                    {
                        "number": 1,
                        "title": "Intro",
                        "heading": "Intro",
                        "word_target": 1000,
                    }
                ],
            }
        )

        mock_instructor = MagicMock()
        mock_instructor.chat.completions.create.return_value = complete_plan
        client = ModelClient(
            client=mock_instructor, model="test-model", model_spec="openai:test-model"
        )

        # Patch _retry_with_backoff to just call the fn
        monkeypatch.setattr(
            "src.pipeline_support._retry_with_backoff", lambda fn, **kw: fn()
        )

        result = _structured_call(client, "Plan prompt", EssayPlan)

        assert len(result.sections) == 1
        mock_instructor.chat.completions.create.assert_called_once()
        call_kwargs = mock_instructor.chat.completions.create.call_args
        assert call_kwargs.kwargs["response_model"] is EssayPlan
        assert call_kwargs.kwargs["model"] == "test-model"

    def test_async_structured_call_uses_instructor(self, monkeypatch):
        """Verify _async_structured_call delegates to async Instructor."""
        from src.pipeline_support import _async_structured_call
        from src.schemas import EssayPlan
        from src.agent import AsyncModelClient

        complete_plan = EssayPlan.model_validate(
            {
                "title": "Test",
                "thesis": "Test thesis",
                "research_queries": ["query"],
                "total_word_target": 1000,
                "sections": [
                    {
                        "number": 1,
                        "title": "Intro",
                        "heading": "Intro",
                        "word_target": 1000,
                    }
                ],
            }
        )

        mock_instructor = MagicMock()

        async def fake_create(**kwargs):
            return complete_plan

        mock_instructor.chat.completions.create = fake_create
        client = AsyncModelClient(
            client=mock_instructor, model="test-model", model_spec="openai:test-model"
        )

        # Patch _retry_with_backoff to handle async
        async def fake_retry(fn, *, is_async=False):
            return await fn()

        monkeypatch.setattr(
            "src.pipeline_support._retry_with_backoff", fake_retry
        )

        result = asyncio.run(_async_structured_call(client, "Plan prompt", EssayPlan))

        assert len(result.sections) == 1


# ── web_fetcher HTML stripping ────────────────────────────────────────────


class TestHtmlToText:
    def test_strips_tags(self):
        from src.tools.web_fetcher import _html_to_text

        assert _html_to_text("<p>Hello <b>world</b></p>") == "Hello world"

    def test_skips_script_and_style(self):
        from src.tools.web_fetcher import _html_to_text

        html = "<div>before</div><script>var x = 1;</script><div>after</div>"
        text = _html_to_text(html)
        assert "var x" not in text
        assert "before" in text
        assert "after" in text

    def test_collapses_whitespace(self):
        from src.tools.web_fetcher import _html_to_text

        html = "<p>a</p>" + "<br>" * 10 + "<p>b</p>"
        text = _html_to_text(html)
        assert "\n\n\n" not in text

    def test_handles_attributes_with_angle_brackets(self):
        from src.tools.web_fetcher import _html_to_text

        html = '<div title="a > b">content</div>'
        text = _html_to_text(html)
        assert "content" in text


# ── search error response ────────────────────────────────────────────────


class TestSearchErrorResponse:
    def test_format(self):
        import json

        from src.tools._http import search_error_response

        result = json.loads(
            search_error_response("crossref", "test query", ValueError("oops"))
        )
        assert result["error"] == "request_failed"
        assert result["source"] == "crossref"
        assert result["query"] == "test query"
        assert "oops" in result["message"]


class _FakeResponse:
    def __init__(self, status_code=200, text="ok", headers=None):
        self.status_code = status_code
        self.text = text
        self.headers = headers or {}
        self.content = text.encode("utf-8")

    @property
    def is_error(self):
        return self.status_code >= 400

    def raise_for_status(self):
        if self.is_error:
            request = httpx.Request("GET", "https://example.com")
            response = httpx.Response(self.status_code, request=request)
            raise httpx.HTTPStatusError("boom", request=request, response=response)


class TestSharedHttp:
    def test_http_get_retries_request_errors(self, monkeypatch):
        from src.tools import _http

        calls = {"count": 0}

        class FakeClient:
            def get(self, *args, **kwargs):
                calls["count"] += 1
                if calls["count"] == 1:
                    raise httpx.ConnectError(
                        "nope", request=httpx.Request("GET", "https://example.com")
                    )
                return _FakeResponse()

        monkeypatch.setattr(_http, "get_http_client", lambda: FakeClient())
        monkeypatch.setattr(_http.time, "sleep", lambda _: None)

        response = _http.http_get(
            "https://example.com", max_retries=1, request_name="test"
        )

        assert response.status_code == 200
        assert calls["count"] == 2


class TestResearchConcurrency:
    def test_query_worker_count_is_bounded(self):
        from src.tools.research_sources import _query_worker_count

        assert _query_worker_count(0) == 1
        assert _query_worker_count(1) == 1
        assert _query_worker_count(2) == 2
        assert _query_worker_count(10) == 3

    def test_run_queries_parallelizes_but_preserves_query_order(self, monkeypatch):
        from src.tools import research_sources

        active = 0
        max_active = 0
        lock = threading.Lock()

        def fake_search_one_query(query, max_per_api):
            nonlocal active, max_active
            with lock:
                active += 1
                max_active = max(max_active, active)
            try:
                if query == "q1":
                    time.sleep(0.03)
                else:
                    time.sleep(0.01)
            finally:
                with lock:
                    active -= 1

            return ([{"title": query}], {"crossref": {"query": query}})

        monkeypatch.setattr(
            research_sources, "_search_one_query", fake_search_one_query
        )

        all_results, all_raw = research_sources._run_queries(["q1", "q2", "q3"], 2)

        assert max_active > 1
        assert [item["title"] for item in all_results] == ["q1", "q2", "q3"]
        assert [item["query"] for item in all_raw] == ["q1", "q2", "q3"]

    def test_http_get_retries_retryable_statuses(self, monkeypatch):
        from src.tools import _http

        responses = [_FakeResponse(status_code=503), _FakeResponse(status_code=200)]

        class FakeClient:
            def get(self, *args, **kwargs):
                return responses.pop(0)

        monkeypatch.setattr(_http, "get_http_client", lambda: FakeClient())
        monkeypatch.setattr(_http.time, "sleep", lambda _: None)

        response = _http.http_get(
            "https://example.com", max_retries=1, request_name="test"
        )

        assert response.status_code == 200


class TestConfigBackedBehavior:
    def test_citation_rank_sorts_higher_citations_first(self):
        from src.tools.research_sources import _build_registry

        raw_results = [
            {
                "title": "Low citations",
                "authors": ["A Smith"],
                "year": 2024,
                "abstract": "",
                "doi": "10.1/low",
                "url": "https://example.com/low",
                "pdf_url": "https://example.com/low.pdf",
                "source_type": "",
                "citation_count": 5,
            },
            {
                "title": "High citations",
                "authors": ["B Jones"],
                "year": 2024,
                "abstract": "",
                "doi": "10.1/high",
                "url": "https://example.com/high",
                "pdf_url": "https://example.com/high.pdf",
                "source_type": "",
                "citation_count": 500,
            },
        ]
        registry = _build_registry(raw_results, 10)
        ids = list(registry.keys())
        assert len(ids) == 2
        # Higher citations should come first (both have same accessibility)
        assert registry[ids[0]]["title"] == "High citations"
        assert registry[ids[1]]["title"] == "Low citations"

    def test_build_registry_ranks_by_citations_then_accessibility(self):
        from src.tools.research_sources import _build_registry

        raw_results = [
            {
                "title": "DOI only paper",
                "authors": ["Alice Smith"],
                "year": 2024,
                "abstract": "Some abstract",
                "doi": "10.1/a",
                "url": "",
                "pdf_url": "",
                "source_type": "journal-article",
                "citation_count": 100,
            },
            {
                "title": "OA PDF paper",
                "authors": ["Bob Jones"],
                "year": 2024,
                "abstract": "Another abstract",
                "doi": "10.1/b",
                "url": "https://example.com/b",
                "pdf_url": "https://example.com/b.pdf",
                "source_type": "journal-article",
                "citation_count": 10,
            },
        ]

        registry = _build_registry(raw_results, 10)
        ids = list(registry.keys())
        assert len(ids) == 2
        # Higher citations should rank first; accessibility is tiebreaker
        assert registry[ids[0]]["title"] == "DOI only paper"

    def test_rendered_review_prompt_uses_configured_tolerance(self):
        from src.rendering import render_prompt
        from src.pipeline import Section

        prompt = render_prompt(
            "section_review.j2",
            section=Section(number=1, title="Intro", heading="Intro", word_target=100),
            full_essay="Body",
            section_words=96,
            tolerance_ratio=0.05,
            tolerance_percent=5,
            tolerance_ratio_over=0.20,
            tolerance_percent_over=20,
            language="English",
        )

        assert "at least" in prompt.lower()


class TestValidationQuestionSuggestedIndex:
    def test_clamps_suggested_option_index_to_valid_range(self):
        from src.schemas import ValidationQuestion

        q = ValidationQuestion(
            question="Q?",
            options=["a", "b", "c"],
            suggested_option_index=99,
        )
        assert q.suggested_option_index == 2

        q2 = ValidationQuestion(
            question="Q?",
            options=["only"],
            suggested_option_index=-3,
        )
        assert q2.suggested_option_index == 0


class TestValidationClarifications:
    def test_parse_validation_answers_maps_letter_choices(self):
        from src.runner import _parse_validation_answers
        from src.schemas import ValidationQuestion

        questions = [
            ValidationQuestion(
                question="Choose scope",
                options=["Macro", "Micro", "Mixed"],
            ),
            ValidationQuestion(
                question="Need case study",
                options=["Yes", "No"],
            ),
        ]

        clarifications = _parse_validation_answers(questions, "1. b, 2. a")

        assert len(clarifications) == 2
        assert clarifications[0].answer == "Micro"
        assert clarifications[1].answer == "Yes"

    def test_parse_validation_answers_allows_single_question_freeform(self):
        from src.runner import _parse_validation_answers
        from src.schemas import ValidationQuestion

        questions = [
            ValidationQuestion(
                question="Clarify the focus",
                options=["Option A", "Option B"],
            )
        ]

        clarifications = _parse_validation_answers(
            questions, "Focus on public policy implications"
        )

        assert len(clarifications) == 1
        assert clarifications[0].answer == "Focus on public policy implications"

    def test_handle_questions_persists_structured_clarifications(
        self, tmp_path, monkeypatch
    ):
        from src.runner import _handle_questions
        from src.schemas import AssignmentBrief, ValidationQuestion

        brief_path = tmp_path / "brief" / "assignment.json"
        brief_path.parent.mkdir(parents=True)
        brief = AssignmentBrief(topic="Topic", description="Desc")
        brief_path.write_text(brief.model_dump_json(), encoding="utf-8")

        questions = [
            ValidationQuestion(
                question="Choose scope",
                options=["Macro", "Micro"],
            ),
            ValidationQuestion(
                question="Need case study",
                options=["Yes", "No"],
            ),
        ]
        monkeypatch.setattr("builtins.input", lambda _prompt="": "1. b, 2. a")

        _handle_questions(questions, tmp_path)

        saved = AssignmentBrief.model_validate_json(
            brief_path.read_text(encoding="utf-8")
        )

        assert saved.clarifications is not None
        assert len(saved.clarifications) == 2
        assert saved.clarifications[0].question == "Choose scope"
        assert saved.clarifications[0].answer == "Micro"
        assert saved.clarifications[1].answer == "Yes"


class TestPricing:
    def test_calc_cost_known_model(self):
        from src.runner import _calc_cost

        cost = _calc_cost(
            "gemini-2.5-flash", input_tokens=1_000_000, output_tokens=1_000_000
        )
        assert cost > 0

    def test_calc_cost_unknown_model_returns_zero(self):
        from src.runner import _calc_cost

        cost = _calc_cost("nonexistent-model-xyz", input_tokens=1000, output_tokens=100)
        assert cost == 0.0

    def test_calc_cost_includes_thinking(self):
        from src.runner import _calc_cost

        cost_no_think = _calc_cost("gpt-4o", input_tokens=1000, output_tokens=100)
        cost_with_think = _calc_cost(
            "gpt-4o", input_tokens=1000, output_tokens=100, thinking_tokens=500
        )
        assert cost_with_think > cost_no_think


# ── intake classify ──────────────────────────────────────────────────────


class TestClassify:
    @pytest.mark.parametrize(
        "filename,expected",
        [
            ("doc.pdf", "pdf"),
            ("doc.docx", "docx"),
            ("doc.pptx", "pptx"),
            ("image.png", "image"),
            ("image.jpg", "image"),
            ("notes.txt", "text"),
            ("notes.md", "text"),
            ("data.csv", "text"),
            ("file.xyz", "unsupported"),
        ],
    )
    def test_classify(self, filename, expected):
        from src.intake import _classify

        assert _classify(Path(filename)) == expected


# ── intake base64 helper ─────────────────────────────────────────────────


class TestMakeImageBlock:
    def test_produces_valid_block(self):
        from src.intake import _make_image_block

        block = _make_image_block(b"fake-png-data", "image/png")
        assert block["type"] == "image_url"
        url = block["image_url"]["url"]
        assert url.startswith("data:image/png;base64,")
        # Verify the base64 decodes back
        encoded = url.split(",", 1)[1]
        assert base64.standard_b64decode(encoded) == b"fake-png-data"


class TestBuildExtractedText:
    def test_builds_extracted_text_with_prompt_and_warnings(self):
        from src.intake import InputFile, build_extracted_text

        files = [
            InputFile(Path("topic.txt"), "text", text="Essay topic"),
            InputFile(Path("scan.pdf"), "pdf", image_blocks=[{"type": "image_url"}]),
            InputFile(
                Path("legacy.doc"),
                "unsupported",
                warning="Old Word binary format — save as .docx first",
            ),
        ]

        extracted = build_extracted_text(files, extra_prompt="Focus on economics")

        assert "### File: topic.txt" in extracted
        assert "Essay topic" in extracted
        assert "### Image: scan.pdf" in extracted
        assert "text extraction was sparse" in extracted
        assert "## Warnings" in extracted
        assert "legacy.doc" in extracted
        assert "## Additional Instructions" in extracted
        assert "Focus on economics" in extracted


# ── docx extraction dedup ────────────────────────────────────────────────


class TestExtractDocxText:
    def test_extracts_headings_and_body(self):
        from src.tools.docx_reader import extract_docx_text

        from docx import Document

        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Body text here.")
        doc.add_heading("Sub", level=2)

        result = extract_docx_text(doc)
        assert "# Title" in result
        assert "Body text here." in result
        assert "## Sub" in result


# ── rendering cache ──────────────────────────────────────────────────────


class TestRendering:
    def test_render_prompt_returns_string(self):
        from src.rendering import render_prompt

        # Test new per-task templates
        result = render_prompt(
            "intake.j2", extracted_text="Test content", extra_prompt=None
        )
        assert isinstance(result, str)
        assert len(result) > 0

        result = render_prompt("validate.j2", brief_json='{"topic": "test"}')
        assert isinstance(result, str)
        assert len(result) > 0

    def test_cached_env_is_same_object(self):
        from src.rendering import _get_env

        _get_env.cache_clear()
        env1 = _get_env("/tmp/dummy")
        env2 = _get_env("/tmp/dummy")
        assert env1 is env2


# ── selected source filtering ─────────────────────────────────────────────


class TestSelectedSourceNotes:
    def test_uses_selected_accessible_notes_when_available(self, tmp_path):
        from src.pipeline import _load_selected_source_notes
        from src.schemas import SourceNote

        notes_dir = tmp_path / "sources" / "notes"
        notes_dir.mkdir(parents=True)

        note_a = SourceNote(source_id="alpha2024", is_accessible=True, title="A")
        note_b = SourceNote(source_id="beta2024", is_accessible=True, title="B")
        (notes_dir / "alpha2024.json").write_text(
            note_a.model_dump_json(), encoding="utf-8"
        )
        (notes_dir / "beta2024.json").write_text(
            note_b.model_dump_json(), encoding="utf-8"
        )

        (tmp_path / "sources" / "selected.json").write_text(
            json.dumps({"beta2024": {"title": "B"}}), encoding="utf-8"
        )

        notes = _load_selected_source_notes(tmp_path)
        assert [note.source_id for note in notes] == ["beta2024"]

    def test_falls_back_to_all_accessible_notes_when_selection_is_unusable(
        self, tmp_path, caplog
    ):
        from src.pipeline import _load_selected_source_notes
        from src.schemas import SourceNote

        notes_dir = tmp_path / "sources" / "notes"
        notes_dir.mkdir(parents=True)

        note_a = SourceNote(source_id="alpha2024", is_accessible=True, title="A")
        note_b = SourceNote(source_id="beta2024", is_accessible=True, title="B")
        (notes_dir / "alpha2024.json").write_text(
            note_a.model_dump_json(), encoding="utf-8"
        )
        (notes_dir / "beta2024.json").write_text(
            note_b.model_dump_json(), encoding="utf-8"
        )

        (tmp_path / "sources" / "selected.json").write_text(
            json.dumps({"missing2024": {"title": "Missing"}}), encoding="utf-8"
        )

        with caplog.at_level(logging.WARNING):
            notes = _load_selected_source_notes(tmp_path)

        assert [note.source_id for note in notes] == ["alpha2024", "beta2024"]
        assert "Selected sources had no accessible notes" in caplog.text

    def test_source_read_candidates_includes_all_api_sources(self):
        from src.pipeline import _source_read_candidates

        registry = {
            f"s{i}": {"title": f"Source {i}", "url": f"https://example.com/{i}"}
            for i in range(1, 21)
        }
        # Add one without URL — should be excluded
        registry["s_no_url"] = {"title": "No URL source"}

        candidates = _source_read_candidates(registry, target_sources=8)

        assert len(candidates) == 20
        assert all(sid != "s_no_url" for sid, _ in candidates)


class TestSourceTargetScaling:
    def test_compute_max_sources_log_scaling(self):
        from config.schemas import EssayWriterConfig
        from src.pipeline import _compute_max_sources, _suggested_sources

        cfg = EssayWriterConfig()
        target, fetch = _compute_max_sources(24000, cfg, None)
        expected = _suggested_sources(24000, cfg.search.sources_per_1k_words)
        assert target == max(cfg.search.min_sources, expected)
        assert fetch == int(target * cfg.search.overfetch_multiplier)
        # Log scaling should produce fewer sources than the old linear formula
        old_linear = 24 * cfg.search.sources_per_1k_words  # 120
        assert target < old_linear

    def test_suggested_sources_values(self):
        """Spot-check the log-based formula at key word counts."""
        from src.pipeline import _suggested_sources

        assert _suggested_sources(0) == 0
        assert 22 <= _suggested_sources(2000) <= 26
        assert 37 <= _suggested_sources(5000) <= 41
        assert 50 <= _suggested_sources(10000) <= 55
        assert 63 <= _suggested_sources(20000) <= 69
        assert 72 <= _suggested_sources(30000) <= 77

    def test_compute_max_sources_respects_user_floor_above_raw(self):
        from config.schemas import EssayWriterConfig
        from src.pipeline import _compute_max_sources

        cfg = EssayWriterConfig()
        target, fetch = _compute_max_sources(24000, cfg, 130)
        assert target == 130
        assert fetch == int(130 * cfg.search.overfetch_multiplier)

    def test_compute_max_sources_explicit_user_above_raw(self):
        """User min (e.g. 90) above log-based suggestion (~65 for 24k) wins."""
        from config.schemas import EssayWriterConfig
        from src.pipeline import _compute_max_sources

        cfg = EssayWriterConfig()
        target, fetch = _compute_max_sources(24000, cfg, 90)
        assert target == 90
        assert fetch == int(90 * cfg.search.overfetch_multiplier)

    def test_compute_max_sources_explicit_user_below_raw(self):
        """User min below the log-based suggestion still uses user value."""
        from config.schemas import EssayWriterConfig
        from src.pipeline import _compute_max_sources

        cfg = EssayWriterConfig()
        target, fetch = _compute_max_sources(24000, cfg, 30)
        assert target == 30
        assert fetch == int(30 * cfg.search.overfetch_multiplier)


class TestLongEssayContextHelpers:
    def test_prior_section_context_uses_recent_sections_only(self):
        from src.pipeline import Section, _build_prior_sections_context

        sections = [
            (Section(number=1, title="One", heading="One", word_target=100), "intro"),
            (Section(number=2, title="Two", heading="Two", word_target=100), "body a"),
            (
                Section(number=3, title="Three", heading="Three", word_target=100),
                "body b",
            ),
        ]

        context = _build_prior_sections_context(sections, max_sections=2)

        assert "intro" not in context
        assert "body a" in context
        assert "body b" in context

    def test_review_context_uses_only_adjacent_sections(self):
        from src.pipeline import Section, _build_review_context

        sections = [
            Section(number=1, title="One", heading="One", word_target=100),
            Section(number=2, title="Two", heading="Two", word_target=100),
            Section(number=3, title="Three", heading="Three", word_target=100),
            Section(number=4, title="Four", heading="Four", word_target=100),
            Section(number=5, title="Five", heading="Five", word_target=100),
        ]
        section_texts = {
            2: "section two",
            3: "section three",
            4: "section four",
        }

        context = _build_review_context(sections[2], sections, section_texts)

        assert "section two" in context
        assert "section three" in context
        assert "section four" in context
        assert "SECTION TO REVIEW: START" in context
        assert "SECTION TO REVIEW: END" in context
        assert "section one" not in context
        assert "section five" not in context


# ── docx_builder table support ────────────────────────────────────────


class TestDocxTableParsing:
    """Tests for markdown table → docx table conversion."""

    def test_simple_table(self):
        from src.tools.docx_builder import build_document

        md = (
            "Some text before.\n"
            "\n"
            "| Name | Age |\n"
            "|------|-----|\n"
            "| Alice | 30 |\n"
            "| Bob | 25 |\n"
            "\n"
            "Some text after."
        )
        doc = build_document(md, {"title": "Test"})
        tables = doc.tables
        assert len(tables) == 1
        table = tables[0]
        # Header row + 2 data rows
        assert len(table.rows) == 3
        assert len(table.columns) == 2
        assert table.rows[0].cells[0].text == "Name"
        assert table.rows[0].cells[1].text == "Age"
        assert table.rows[1].cells[0].text == "Alice"
        assert table.rows[2].cells[1].text == "25"

    def test_table_with_inline_formatting(self):
        from src.tools.docx_builder import build_document

        md = "| Header |\n|--------|\n| **bold** and *italic* |\n"
        doc = build_document(md, {"title": "Test"})
        tables = doc.tables
        assert len(tables) == 1
        # Cell should contain the text (formatting applied via runs)
        cell_text = tables[0].rows[1].cells[0].text
        assert "bold" in cell_text
        assert "italic" in cell_text

    def test_three_column_table(self):
        from src.tools.docx_builder import build_document

        md = (
            "| A | B | C |\n"
            "|---|---|---|\n"
            "| 1 | 2 | 3 |\n"
            "| 4 | 5 | 6 |\n"
            "| 7 | 8 | 9 |\n"
        )
        doc = build_document(md, {"title": "Test"})
        table = doc.tables[0]
        assert len(table.columns) == 3
        assert len(table.rows) == 4  # 1 header + 3 data
        assert table.rows[3].cells[2].text == "9"

    def test_no_table_without_separator(self):
        """Pipe lines without a separator row should NOT be parsed as a table."""
        from src.tools.docx_builder import build_document

        md = "| Not a table |\n| Just pipes |\n"
        doc = build_document(md, {"title": "Test"})
        assert len(doc.tables) == 0

    def test_table_between_paragraphs(self):
        """Table should not swallow surrounding paragraphs."""
        from src.tools.docx_builder import build_document

        md = "Paragraph before.\n\n| X |\n|---|\n| 1 |\n\nParagraph after."
        doc = build_document(md, {"title": "Test"})
        assert len(doc.tables) == 1
        # Check that both paragraphs exist in the document text
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Paragraph before." in full_text
        assert "Paragraph after." in full_text

    def test_header_bold(self):
        """Header cells should be bold."""
        from src.tools.docx_builder import build_document

        md = "| Col |\n|-----|\n| val |\n"
        doc = build_document(md, {"title": "Test"})
        header_cell = doc.tables[0].rows[0].cells[0]
        runs = header_cell.paragraphs[0].runs
        assert any(r.bold for r in runs)


# ── docx_builder heading & citation fixes ─────────────────────────────


class TestHeadingAsterisks:
    """Headings should have markdown bold/italic markers stripped."""

    def test_strips_double_asterisks(self):
        from src.tools.docx_builder import build_document

        md = "## **Bold Heading**\n\nSome text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert any("Bold Heading" in h and "**" not in h for h in headings)

    def test_strips_single_asterisks(self):
        from src.tools.docx_builder import build_document

        md = "## *Italic Heading*\n\nSome text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert any("Italic Heading" in h and "*" not in h for h in headings)

    def test_strips_triple_asterisks(self):
        from src.tools.docx_builder import build_document

        md = "## ***Bold Italic Heading***\n\nSome text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert any("Bold Italic Heading" in h and "*" not in h for h in headings)


class TestEmptyAuthorCitations:
    """Sources with empty/blank authors should fall back to title-based citation."""

    def test_empty_author_list_inline(self):
        from src.tools.docx_builder import _format_apa_inline

        source = {"authors": [], "title": "Some Title", "year": 2020}
        result = _format_apa_inline(source, None)
        assert "Some Title" in result
        assert "& " not in result

    def test_blank_authors_inline(self):
        from src.tools.docx_builder import _format_apa_inline

        source = {"authors": ["", "  "], "title": "My Paper", "year": 2017}
        result = _format_apa_inline(source, None)
        assert "My Paper" in result
        assert "( & " not in result

    def test_blank_authors_bib_entry(self):
        from src.tools.docx_builder import _format_bib_entry

        source = {"authors": ["", ""], "title": "A Title", "year": 2020}
        result = _format_bib_entry(source)
        assert result.startswith("Unknown (2020)")
        assert ", ," not in result


class TestH1TitleSkipping:
    """First H1 should be skipped since it duplicates the cover page title."""

    def test_first_h1_skipped(self):
        from src.tools.docx_builder import build_document

        md = "# Essay Title\n\n## Section One\n\nSome text."
        doc = build_document(md, {"title": "Essay Title"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert "Essay Title" not in headings
        assert "Section One" in headings

    def test_second_h1_kept(self):
        from src.tools.docx_builder import build_document

        md = "# First Title\n\n## Section\n\nText.\n\n# Second H1\n\nMore text."
        doc = build_document(md, {"title": "Test"})
        headings = [
            p.text for p in doc.paragraphs if p.style.name.startswith("Heading")
        ]
        assert "First Title" not in headings
        assert "Second H1" in headings


class TestTableCellAlignment:
    """Table cells should be left-aligned, not justified."""

    def test_cells_left_aligned(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        from src.tools.docx_builder import build_document

        md = "| A | B |\n|---|---|\n| 1 | 2 |\n"
        doc = build_document(md, {"title": "Test"})
        table = doc.tables[0]
        for row in table.rows:
            for cell in row.cells:
                assert cell.paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
